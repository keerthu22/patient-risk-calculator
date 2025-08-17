import streamlit as st
import pandas as pd
import joblib
from model.preprocess import preprocess
from model.featureEngineer import featureEngineer
import joblib
from docx import Document
from docx.shared import Inches
import io
from model.keys import api_key
from model.PreventiveCare import PreventiveAdvisor

advisor = PreventiveAdvisor(api_key)

def generate_patient_report(df, summ, advisor):
    patient_docs = {}
    grouped = df.groupby("MemberID")

    for member_id, group in grouped:
        doc = Document()
        doc.add_heading("🧠 Patient Risk Report", level=0)
        doc.add_paragraph(f"Member ID: {member_id}")
        doc.add_paragraph(f"Age: {group['Age'].iloc[0]}")
        doc.add_paragraph(f"Gender: {group['Gender'].iloc[0]}")
        doc.add_paragraph(f"Total Claims: {len(group)}")

        doc.add_heading("📄 Claim Details", level=1)

        adv_set = set()
        prev_care = ""

        for _, row in group.iterrows():
            doc.add_paragraph(
                f"• Claim ID: {row['ClaimID']}\n"
                f"  - Diagnosis Code: {row['DiagnosisCode']} ({row['DiseaseName']})\n"
                f"  - Procedure Code: {row['ProcedureCode']}\n"
                f"  - Disease Criticality: {row['DiseaseCriticality']}\n"
                f"  - Amount Billed: ${row['AmountBilled']}\n",
                style='List Bullet'
            )

            # adv = row['PreventiveCareAdvice']
            # if adv and adv not in adv_set:
            #     adv_set.add(adv)

        # Summary row
        row = summ[summ['MemberID'] == member_id].iloc[0]
        # prev_care_str = "; ".join(adv_set) if adv_set else "Not available"

        doc.add_paragraph(
            f"  Avg Criticality: {row['AvgCriticality']}\n"
            f"  Max Criticality: {row['MaxCriticality']}\n"
            f"  Chronic Count: {row['ChronicCount']}\n"
            f"  Risk Score: {row['PredictedRisk']}\n"
            f"  Recommendation: {generate_recommendation(row['PredictedRisk'])}\n"
            # f"  Preventive Care Advice: {prev_care_str}"
        )

        # --- Extract plain text from docx for Gemini ---
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Call Gemini advisor
        geminiResult = advisor.generate_advice(full_text)

        doc.add_paragraph("\nAI-Generated Preventive Care Advice:")
        doc.add_paragraph(geminiResult)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        patient_docs[member_id] = buffer

    return patient_docs


st.set_page_config(page_title="🧠 Patient Risk Analyzer", layout="wide")

# Title and subtitle
st.markdown("<h1 style='text-align: center;'>🧠 Patient Risk Analyzer</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center;'>Predict and monitor health risks using AI-powered analysis of patient claim data.</p>", unsafe_allow_html=True)
st.markdown(
    """
    <style>
        /* Full app background */
        .stApp {
            background: linear-gradient(to bottom right, #f0f4ff, #e0f7fa);
            background-size: cover;
            color:rgb(0,0,0);
        }

        /* Make table backgrounds slightly transparent */
        .stDataFrameContainer {
            background-color: rgba(255, 255, 255, 0.95) !important;
            border-radius: 10px;
            padding: 10px;
        }

        /* Style the file uploader and buttons */
        .stButton>button, .stDownloadButton>button {
            border-radius: 8px;
            background-color: #1976d2;
            color: white;
        }

        .stDownloadButton>button:hover {
            background-color: #125a9c;
        }

        /* Tabs styling */
        .stTabs [role="tablist"] {
            background-color: #dbeafe;
            border-radius: 10px;
        }
    </style>
    """,
    unsafe_allow_html=True
)
st.markdown("---")

# Tabs
tab1, tab2 = st.tabs(["📁 Upload & Analyze", "📊 About & Insights"])

with tab1:
    st.markdown("### 📤 Upload Patient CSV File")
    uploaded_file = st.file_uploader("Upload your patient claims CSV file", type=["csv"])

    if uploaded_file:
        try:
            # Show uploaded file
            raw_df = pd.read_csv(uploaded_file)
            st.markdown("#### 📄 Uploaded Input Data")
            st.dataframe(raw_df, use_container_width=True)

            with st.spinner("🔄 Processing data..."):
                # Preprocess and feature engineer
                cleaned_df = preprocess(raw_df)
                summary = featureEngineer(cleaned_df)[0].reset_index(drop=True)
                report = featureEngineer(cleaned_df)[1].reset_index(drop=True)

                # Load trained model
                model = joblib.load("output/patient_risk_model.pkl")

                # Define features
                features = ['Age', 'Gender', 'NumClaims', 'UniqueDiseases', 'AvgCriticality', 'MaxCriticality', 'ChronicCount', 'TotalAmountBilled', 'DiagnosisCode', 'ProcedureCode']

                # Predict
                summary['PredictedRisk'] = model.predict(summary[features])

                # Recommendation logic
                def generate_recommendation(score):
                    score_dict = {
                        1 : 'Routine Care 🟢',
                        2 : 'Monitor Periodically 🔵',
                        3 : 'Monitor Closely 🟠',
                        4 : 'High Priority 🟣',
                        5 : 'Immediate Attention 🔴'
                    }

                    return score_dict[score]

                summary['Recommendation'] = summary['PredictedRisk'].apply(generate_recommendation)
                
                # Load the encoders
                le_gender = joblib.load("output/le_gender.pkl")
                le_proc = joblib.load("output/le_proc.pkl")
                le_diag = joblib.load("output/le_diag.pkl")

                # Inverse transform the encoded columns
                summary['Gender'] = le_gender.inverse_transform(summary['Gender'])
                summary['DiagnosisCode'] = le_diag.inverse_transform(summary['DiagnosisCode'])
                summary['ProcedureCode'] = le_proc.inverse_transform(summary['ProcedureCode'])

                # Adding the disease Category
                icd_df = pd.read_csv('data/final_datasets/final_ICD.csv')

                code_to_desc = icd_df.set_index('Code')['Desc'].to_dict()
                group_to_cat = icd_df.set_index('Group')['Category'].to_dict()

                def get_disease_name(code):
                    code = code.replace(".", "")
                    if code in code_to_desc:
                        return code_to_desc[code]
                    # 3-char group fallback
                    prefix = code[:3]
                    return group_to_cat.get(prefix, 'Unknown')


                summary['DiseaseName'] = summary['DiagnosisCode'].apply(get_disease_name)

                summary.rename(columns={'DiseaseName':'Category'}, inplace=True)

                # Display final summary
                st.success("✅ Predictions generated successfully!")
                st.markdown("### 📋 Patient Risk Summary")
                st.dataframe(summary[['MemberID', 'Age', 'Gender', 'DiagnosisCode', 'Category', 'ProcedureCode', 'NumClaims', 'UniqueDiseases',
                                         'AvgCriticality', 'MaxCriticality', 'ChronicCount',
                                         'PredictedRisk', 'Recommendation']].reset_index(drop=True),
                             use_container_width=True)

                # Download option
                csv = summary.to_csv(index=False).encode('utf-8')
                st.download_button("📥 Download Results as CSV", data=csv,
                                   file_name="patient_predictions.csv", mime="text/csv")
                
                # Generate reports for each patient
                patient_reports = generate_patient_report(report, summary, advisor)

                st.markdown("### 📄 Download Individual Patient Reports (DOCX)")
                for member_id, doc_buffer in patient_reports.items():
                    st.download_button(
                        label=f"📥 Download Report for Member {member_id}",
                        data=doc_buffer,
                        file_name=f"Patient_Report_{member_id}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        except Exception as e:
            st.error(f"❌ Failed to process: {e}")
    else:
        st.info("Upload your CSV to begin.")

with tab2:
    st.markdown("## 🧾 How It Works")
    st.markdown("""
    - The model takes patient claim data with fields like age, gender, diagnosis, procedures, and billing.
    - It applies advanced feature engineering (chronic disease count, disease diversity, etc.).
    - Predictions are made using a machine learning model trained on past claim data.
    - Risk levels:
        - 🔵 Routine Care – Regular checkups only
        - 🟢 Low Risk – Stable, minor issues
        - 🟠 Medium Risk – Requires monitoring
        - 🟣 High Priority – Elevated risk, proactive care needed
        - 🔴 Immediate Attention – Critical, urgent care required
    """)

    st.markdown("## 🤖 Model Details")
    st.markdown("""
    - **Algorithm:** Gradient Boosting (or similar)
    - **Features used:** Clinical & financial claim features
    - **Custom logic:** Feature extraction based on medical codes
    """)

    st.markdown("## 🔐 Data Privacy")
    st.markdown("All data is processed locally and not stored anywhere. Your privacy is safe!")
